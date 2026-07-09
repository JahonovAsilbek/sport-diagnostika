"""Format renderers — a `ReportDataset` → file bytes. WeasyPrint is imported lazily inside the
PDF renderer only, so the module (and the Excel/Word paths) load without its native libs."""

import io
from html import escape

from apps.reports.models import Report


def render_excel(dataset):
    from openpyxl import Workbook

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Report"
    worksheet.append([dataset.title])
    if dataset.subtitle:
        worksheet.append([dataset.subtitle])
    worksheet.append(dataset.columns)
    for row in dataset.rows:
        worksheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def render_word(dataset):
    from docx import Document

    document = Document()
    document.add_heading(dataset.title, level=1)
    if dataset.subtitle:
        document.add_paragraph(dataset.subtitle)
    table = document.add_table(rows=1, cols=len(dataset.columns))
    header = table.rows[0].cells
    for index, column in enumerate(dataset.columns):
        header[index].text = str(column)
    for row in dataset.rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_report_html(dataset):
    """Pure HTML for the PDF (escaped — a name/value can't inject markup). Unit-testable
    without WeasyPrint."""
    head = f"<h1>{escape(dataset.title)}</h1>"
    if dataset.subtitle:
        head += f"<p>{escape(dataset.subtitle)}</p>"
    header = "".join(f"<th>{escape(str(col))}</th>" for col in dataset.columns)
    body = "".join(
        "<tr>" + "".join(f"<td>{escape(str(value))}</td>" for value in row) + "</tr>"
        for row in dataset.rows
    )
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<style>table{border-collapse:collapse}th,td{border:1px solid #333;padding:4px}</style>"
        f"</head><body>{head}<table><thead><tr>{header}</tr></thead>"
        f"<tbody>{body}</tbody></table></body></html>"
    )


def render_pdf(dataset):
    from weasyprint import HTML  # lazy — needs native libs (pango/cairo), Docker-only (DVPS-9)

    return HTML(string=build_report_html(dataset)).write_pdf()


RENDERERS = {
    Report.Format.EXCEL: render_excel,
    Report.Format.WORD: render_word,
    Report.Format.PDF: render_pdf,
}
